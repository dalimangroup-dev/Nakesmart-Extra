# -*- coding: utf-8 -*-
"""
Build 4 Panduan Lengkap Nakesmart (Pasien, Nakes, Faskes, EO) → DOCX + PDF.

Output:
  Panduan_NAKESMART_untuk_Pasien.docx + .pdf
  Panduan_NAKESMART_untuk_Nakes.docx + .pdf
  Panduan_NAKESMART_untuk_Faskes.docx + .pdf
  Panduan_NAKESMART_untuk_EO.docx + .pdf
  Panduan_Lengkap_NAKESMART.zip (bundle all 8)

Usage: python build_panduan.py
"""
import os
import sys
import io
import zipfile
from pathlib import Path

# Add this folder to sys.path
sys.path.insert(0, str(Path(__file__).parent))

# docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# reportlab for PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, KeepTogether
)
from reportlab.pdfgen import canvas as pdfcanvas

# Content modules
import content_pasien
import content_nakes
import content_faskes
import content_eo

# Brand colors
FOREST_GREEN = "#0F5132"
DEEP_FOREST = "#1B4D3E"
LIME = "#C8FF00"
CHARCOAL = "#1F2937"
LIGHT_BG = "#F8FAFC"
WHITE = "#FFFFFF"

YEAR = "2026"
VERSION = "VERSION 1.0"
COMPANY = "PT GIANNA MEDICAL CENTER"
WEBSITE = "nakesmart.com"


# =====================================================================
# Unicode → ASCII sanitizer
# Helvetica (default reportlab font) can't render emoji + many Unicode
# glyphs (they become black squares). Replace with ASCII equivalents.
# =====================================================================

UNICODE_REPLACEMENTS = {
    # Status/check marks
    '✅': '[OK]',       # ✅
    '❌': '[Tidak]',    # ❌
    '⚠️': '[!]',  # ⚠️
    '⚠': '[!]',        # ⚠ no VS
    '✓': '[v]',        # ✓
    '✗': '[x]',        # ✗
    # Bullets/diamonds
    '◆': '*',          # ◆
    '◇': '*',          # ◇
    '●': '*',          # ●
    '○': '*',          # ○
    # Emoji symbols (replace with text labels)
    '\U0001F691': 'Ambulan: ',     # 🚑
    '\U0001F198': 'Emergency: ',    # 🆘
    '\U0001F692': 'Damkar: ',       # 🚒
    '\U0001F46E': 'Polisi: ',       # 👮
    '☎️': 'Telp: ',       # ☎️
    '☎': 'Telp: ',             # ☎
    '\U0001F4F1': 'Android: ',      # 📱
    '\U0001F34E': 'iPhone: ',       # 🍎
    '\U0001F4BB': 'Desktop: ',      # 💻
    '\U0001F310': 'Web: ',          # 🌐
    '\U0001F512': '',               # 🔒
    '\U0001F6E1️': '',         # 🛡️
    '\U0001F6E1': '',               # 🛡
    '\U0001F389': '*',              # 🎉
    '\U0001F4DD': '',               # 📝
    '\U0001F4D6': '',               # 📖
    '\U0001F4CB': '',               # 📋
    '\U0001F4CA': '',               # 📊
    '\U0001F4C8': '',               # 📈
    '\U0001F4DA': '',               # 📚
    '\U0001F517': '',               # 🔗
    '✨': '',                   # ✨
    # Arrows — Helvetica/standard-14 PDF fonts lack these glyphs -> black box
    '→': '->',
    '←': '<-',
    '↔': '<->',
    '⇒': '=>',
    '⇐': '<=',
    '↑': '^', '↓': 'v',
    # Em-dashes & punctuation (Helvetica has these but be safe)
    # leave em-dash and middot for now (they work in Helvetica)
}

# Variation selectors (zero-width modifiers that break rendering)
VARIATION_SELECTORS = ['︎', '️']


def sanitize_text(text):
    """Replace Unicode emoji & symbols with ASCII equivalents."""
    if not isinstance(text, str):
        return text
    # First, multi-char sequences (with VS16)
    for needle, repl in sorted(UNICODE_REPLACEMENTS.items(), key=lambda kv: -len(kv[0])):
        text = text.replace(needle, repl)
    # Strip lingering variation selectors
    for vs in VARIATION_SELECTORS:
        text = text.replace(vs, '')
    # Strip any remaining char outside BMP that Helvetica likely can't render
    # (heuristic: anything > U+2FFF that's not standard punctuation)
    clean = []
    for ch in text:
        cp = ord(ch)
        if cp < 0x3000:
            clean.append(ch)
        elif ch in ('—', '–', '•', '·'):  # em/en dash, bullet, middot
            clean.append(ch)
        else:
            # Try to keep CJK in case (we don't have any), drop emoji
            if cp >= 0x1F000:  # emoji range
                continue
            clean.append(ch)
    return ''.join(clean)


# =====================================================================
# DOCX Builder
# =====================================================================

def add_page_number(paragraph):
    """Add PAGE field to paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex.replace('#', ''))
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=FOREST_GREEN):
    text = sanitize_text(text)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(28)
    elif level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(color.replace('#', ''))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, italic=False, bold=False, size=11):
    text = sanitize_text(text)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(CHARCOAL.replace('#', ''))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        p.add_run(text)
    return p


def render_paragraph_smart(doc, text):
    """Render text intelligently — detect bullets, headings, normal text."""
    text = sanitize_text(text)
    lines = text.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Bullet detection
        if stripped.startswith(('•', '-', '*')) and not stripped.startswith('---'):
            content = stripped.lstrip('•-* ').strip()
            p = doc.add_paragraph()
            p.style = doc.styles['List Bullet']
            run = p.add_run(content)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor.from_string(CHARCOAL.replace('#', ''))
            p.paragraph_format.space_after = Pt(3)
        elif stripped.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor.from_string(CHARCOAL.replace('#', ''))
            p.paragraph_format.space_after = Pt(3)
        elif stripped.endswith(':') and len(stripped) < 80:
            # Sub-heading
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string(FOREST_GREEN.replace('#', ''))
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        else:
            add_para(doc, stripped, size=10.5)


def build_docx_cover(doc, role, tagline):
    """Cover page."""
    # Add some spacing
    for _ in range(3):
        doc.add_paragraph()

    # Brand mark
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NAKESMART")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor.from_string(FOREST_GREEN.replace('#', ''))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("* * *")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(LIME.replace('#', ''))

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PANDUAN LENGKAP")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(CHARCOAL.replace('#', ''))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"UNTUK {role.upper()}")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(FOREST_GREEN.replace('#', ''))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(tagline)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(CHARCOAL.replace('#', ''))

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"EDISI {YEAR}  ·  {VERSION}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(FOREST_GREEN.replace('#', ''))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COMPANY)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(CHARCOAL.replace('#', ''))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(WEBSITE)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(FOREST_GREEN.replace('#', ''))

    doc.add_page_break()


def build_docx_toc(doc, role, toc_items):
    """Table of contents."""
    add_heading(doc, "DAFTAR ISI", level=1, color=FOREST_GREEN)
    doc.add_paragraph()
    for chapter, sections in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(sanitize_text(chapter))
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(FOREST_GREEN.replace('#', ''))
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        for sect in sections:
            p = doc.add_paragraph(f"   {sanitize_text(sect)}")
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor.from_string(CHARCOAL.replace('#', ''))
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def build_docx_chapter(doc, chapter_title, sections, content_dict):
    """Render a chapter with all its sections."""
    add_heading(doc, chapter_title.upper(), level=1, color=FOREST_GREEN)

    for sect in sections:
        # Extract section key (e.g. "1.1" or "A")
        # Section format: "1.1 Apa Itu..." or "A. Glossarium..."
        parts = sect.split(' ', 1)
        key = parts[0].rstrip('.')
        title = parts[1] if len(parts) > 1 else ''

        add_heading(doc, sect, level=2, color=DEEP_FOREST)

        # Get content
        content = content_dict.get(key, '').strip()
        if content:
            render_paragraph_smart(doc, content)
        else:
            add_para(doc, "[Konten sedang disiapkan]", italic=True, size=10)

    doc.add_page_break()


def build_docx_footer(doc, role):
    """Add footer with company info."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"© {YEAR} {COMPANY}  ·  {WEBSITE}  ·  Panduan untuk {role}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(CHARCOAL.replace('#', ''))


def build_docx(role, tagline, toc_items, content_dict, output_path):
    """Build complete DOCX for one role."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    build_docx_cover(doc, role, tagline)
    build_docx_toc(doc, role, toc_items)

    for chapter, sections in toc_items:
        build_docx_chapter(doc, chapter, sections, content_dict)

    build_docx_footer(doc, role)

    doc.save(output_path)
    print(f"  DOCX saved: {output_path}")


# =====================================================================
# PDF Builder (reportlab)
# =====================================================================

class PageNumberCanvas(pdfcanvas.Canvas):
    """Add page number to each page."""
    def __init__(self, *args, role="", **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []
        self._role = role

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for idx, state in enumerate(self._saved_page_states, 1):
            self.__dict__.update(state)
            self.draw_footer(idx, num_pages)
            super().showPage()
        super().save()

    def draw_footer(self, page_num, total):
        if page_num == 1:
            return  # skip cover
        self.setFont("Helvetica", 8)
        self.setFillColor(HexColor(CHARCOAL))
        self.drawCentredString(
            A4[0] / 2, 1.2 * cm,
            f"Halaman {page_num}  ·  © {YEAR} {COMPANY}  ·  {WEBSITE}  ·  Panduan untuk {self._role}"
        )


def get_pdf_styles():
    """Define paragraph styles for PDF."""
    base = getSampleStyleSheet()
    styles = {
        'cover_brand': ParagraphStyle(
            'cover_brand', parent=base['Title'],
            fontName='Helvetica-Bold', fontSize=40, textColor=HexColor(FOREST_GREEN),
            alignment=TA_CENTER, leading=46, spaceAfter=8,
        ),
        'cover_mark': ParagraphStyle(
            'cover_mark', parent=base['Title'],
            fontName='Helvetica', fontSize=24, textColor=HexColor(LIME),
            alignment=TA_CENTER, spaceAfter=20,
        ),
        'cover_title': ParagraphStyle(
            'cover_title', parent=base['Title'],
            fontName='Helvetica-Bold', fontSize=32, textColor=HexColor(CHARCOAL),
            alignment=TA_CENTER, leading=38, spaceAfter=8,
        ),
        'cover_role': ParagraphStyle(
            'cover_role', parent=base['Title'],
            fontName='Helvetica-Bold', fontSize=32, textColor=HexColor(FOREST_GREEN),
            alignment=TA_CENTER, leading=38, spaceAfter=20,
        ),
        'cover_tagline': ParagraphStyle(
            'cover_tagline', parent=base['Normal'],
            fontName='Helvetica-Oblique', fontSize=13, textColor=HexColor(CHARCOAL),
            alignment=TA_CENTER, spaceAfter=20,
        ),
        'cover_meta': ParagraphStyle(
            'cover_meta', parent=base['Normal'],
            fontName='Helvetica-Bold', fontSize=11, textColor=HexColor(FOREST_GREEN),
            alignment=TA_CENTER, spaceAfter=4,
        ),
        'cover_company': ParagraphStyle(
            'cover_company', parent=base['Normal'],
            fontName='Helvetica', fontSize=10, textColor=HexColor(CHARCOAL),
            alignment=TA_CENTER, spaceAfter=2,
        ),
        'h1': ParagraphStyle(
            'h1', parent=base['Heading1'],
            fontName='Helvetica-Bold', fontSize=20, textColor=HexColor(FOREST_GREEN),
            spaceBefore=18, spaceAfter=12, leading=24,
        ),
        'h2': ParagraphStyle(
            'h2', parent=base['Heading2'],
            fontName='Helvetica-Bold', fontSize=14, textColor=HexColor(DEEP_FOREST),
            spaceBefore=14, spaceAfter=6, leading=18,
        ),
        'h3': ParagraphStyle(
            'h3', parent=base['Heading3'],
            fontName='Helvetica-Bold', fontSize=11, textColor=HexColor(FOREST_GREEN),
            spaceBefore=8, spaceAfter=4, leading=14,
        ),
        'body': ParagraphStyle(
            'body', parent=base['Normal'],
            fontName='Helvetica', fontSize=10, textColor=HexColor(CHARCOAL),
            alignment=TA_JUSTIFY, leading=13, spaceAfter=4,
        ),
        'bullet': ParagraphStyle(
            'bullet', parent=base['Normal'],
            fontName='Helvetica', fontSize=9.5, textColor=HexColor(CHARCOAL),
            leftIndent=14, spaceAfter=2, leading=12.5, bulletIndent=4,
        ),
        'toc_chapter': ParagraphStyle(
            'toc_chapter', parent=base['Normal'],
            fontName='Helvetica-Bold', fontSize=11.5, textColor=HexColor(FOREST_GREEN),
            spaceBefore=6, spaceAfter=2, leading=14,
        ),
        'toc_section': ParagraphStyle(
            'toc_section', parent=base['Normal'],
            fontName='Helvetica', fontSize=10, textColor=HexColor(CHARCOAL),
            leftIndent=14, spaceAfter=1, leading=12,
        ),
    }
    return styles


def escape_xml(text):
    text = sanitize_text(text)
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def render_pdf_content(text, styles, story):
    """Render content text intelligently for PDF."""
    text = sanitize_text(text)
    lines = text.split('\n')
    bullet_buffer = []

    def flush_bullets():
        nonlocal bullet_buffer
        for b in bullet_buffer:
            story.append(Paragraph("•&nbsp;&nbsp;" + escape_xml(b), styles['bullet']))
        bullet_buffer = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_bullets()
            continue

        if stripped.startswith(('•', '-', '*')) and not stripped.startswith('---'):
            content = stripped.lstrip('•-* ').strip()
            bullet_buffer.append(content)
        elif (stripped.endswith(':') and len(stripped) < 80
              and not stripped.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.'))):
            flush_bullets()
            story.append(Paragraph(escape_xml(stripped), styles['h3']))
        else:
            flush_bullets()
            story.append(Paragraph(escape_xml(stripped), styles['body']))

    flush_bullets()


def build_pdf(role, tagline, toc_items, content_dict, output_path):
    """Build complete PDF for one role."""
    styles = get_pdf_styles()
    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2.2 * cm, rightMargin=2.2 * cm,
    )

    story = []

    # Cover
    story.append(Spacer(1, 3 * cm))
    story.append(Paragraph("NAKESMART", styles['cover_brand']))
    story.append(Paragraph("* * *", styles['cover_mark']))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("PANDUAN LENGKAP", styles['cover_title']))
    story.append(Paragraph(f"UNTUK {role.upper()}", styles['cover_role']))
    story.append(Paragraph(tagline, styles['cover_tagline']))
    story.append(Spacer(1, 4 * cm))
    story.append(Paragraph(f"EDISI {YEAR}  ·  {VERSION}", styles['cover_meta']))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(COMPANY, styles['cover_company']))
    story.append(Paragraph(WEBSITE, styles['cover_company']))
    story.append(PageBreak())

    # TOC
    story.append(Paragraph("DAFTAR ISI", styles['h1']))
    story.append(Spacer(1, 0.3 * cm))
    for chapter, sections in toc_items:
        story.append(Paragraph(escape_xml(chapter), styles['toc_chapter']))
        for sect in sections:
            story.append(Paragraph(escape_xml(sect), styles['toc_section']))
    story.append(PageBreak())

    # Chapters
    for chapter, sections in toc_items:
        story.append(Paragraph(escape_xml(chapter.upper()), styles['h1']))
        for sect in sections:
            parts = sect.split(' ', 1)
            key = parts[0].rstrip('.')
            story.append(Paragraph(escape_xml(sect), styles['h2']))
            content = content_dict.get(key, '').strip()
            if content:
                render_pdf_content(content, styles, story)
            else:
                story.append(Paragraph("<i>[Konten sedang disiapkan]</i>", styles['body']))
        story.append(PageBreak())

    # Build with custom canvas for page numbers
    def canvas_maker(*args, **kwargs):
        return PageNumberCanvas(*args, role=role, **kwargs)

    doc.build(story, canvasmaker=canvas_maker)
    print(f"  PDF saved:  {output_path}")


# =====================================================================
# Build all
# =====================================================================

def build_all():
    out_dir = Path(__file__).parent
    print(f"Building panduan to: {out_dir}\n")

    roles = [
        ("Pasien", content_pasien),
        ("Nakes", content_nakes),
        ("Faskes", content_faskes),
        ("EO", content_eo),
    ]

    output_files = []

    for role_name, module in roles:
        print(f"=== Building {role_name} ===")
        base = f"Panduan_NAKESMART_untuk_{role_name}"
        docx_path = out_dir / f"{base}.docx"
        pdf_path = out_dir / f"{base}.pdf"

        try:
            build_docx(module.ROLE, module.TAGLINE, module.TOC, module.CONTENT, str(docx_path))
            output_files.append(docx_path)
        except Exception as e:
            print(f"  ERROR DOCX {role_name}: {e}")

        try:
            build_pdf(module.ROLE, module.TAGLINE, module.TOC, module.CONTENT, str(pdf_path))
            output_files.append(pdf_path)
        except Exception as e:
            print(f"  ERROR PDF {role_name}: {e}")

        print()

    # Bundle into zip
    zip_path = out_dir / "Panduan_Lengkap_NAKESMART.zip"
    print(f"Creating zip: {zip_path}")
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for f in output_files:
            zf.write(f, f.name)
            print(f"  + {f.name}")

    print(f"\n✅ Done! Zip created: {zip_path}")
    print(f"   Total files: {len(output_files)}")


if __name__ == '__main__':
    build_all()
