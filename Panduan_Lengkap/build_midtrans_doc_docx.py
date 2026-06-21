# -*- coding: utf-8 -*-
"""
Build Dokumen Midtrans Nakesmart dalam format DOCX (editable).

Output: Dokumen_Midtrans_Nakesmart.docx

Mirror dari build_midtrans_doc.py (PDF version) tetapi dengan:
- Placeholder area untuk screenshot asli (bisa di-paste/replace di Word)
- Format Word standar (table, heading, list)
- Mudah di-edit untuk tambah detail spesifik
"""
import os
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand colors
FOREST_GREEN = "0F5132"
DEEP_FOREST = "1B4D3E"
LIME = "C8FF00"
CHARCOAL = "1F2937"
LIGHT_BG = "F8FAFC"
WHITE = "FFFFFF"
MIDTRANS_BLUE = "1F60E3"
GRAY = "94A3B8"
LIGHTGRAY = "E2E8F0"
GREEN_SUCCESS = "DCFCE7"

YEAR = "2026"
COMPANY = "PT GIANNA MEDICAL CENTER"
WEBSITE = "nakesmart.com"


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell, color_hex=LIGHTGRAY, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)
    tc_pr.append(tcBorders)


def add_heading(doc, text, level=1, color=FOREST_GREEN, after=8, before=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    sizes = {0: 28, 1: 20, 2: 14, 3: 11}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def add_para(doc, text, size=10.5, bold=False, italic=False, color=CHARCOAL, align=None):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(CHARCOAL)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bullet_rich(doc, runs, size=10):
    """Add a bullet with multiple runs (bold + normal)."""
    p = doc.add_paragraph(style='List Bullet')
    for (text, bold) in runs:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(CHARCOAL)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_step_header(doc, step_num, title):
    """Step number + title in 2-column table."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    cell_step = table.rows[0].cells[0]
    cell_title = table.rows[0].cells[1]
    cell_step.width = Cm(3)
    cell_title.width = Cm(13)

    set_cell_bg(cell_step, FOREST_GREEN)
    set_cell_bg(cell_title, LIGHT_BG)
    set_cell_border(cell_step, FOREST_GREEN)
    set_cell_border(cell_title, LIGHTGRAY)

    p1 = cell_step.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"STEP {step_num}")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor.from_string(WHITE)

    p2 = cell_title.paragraphs[0]
    r2 = p2.add_run(title)
    r2.bold = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor.from_string(FOREST_GREEN)
    cell_title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add spacer after
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_screenshot_placeholder(doc, screen_title, url, description, mockup_lines):
    """Create a screenshot placeholder area with description below."""
    # Title box
    table = doc.add_table(rows=2, cols=1)
    cell_title = table.rows[0].cells[0]
    cell_mock = table.rows[1].cells[0]

    # Title row
    set_cell_bg(cell_title, FOREST_GREEN)
    set_cell_border(cell_title, FOREST_GREEN)
    p = cell_title.paragraphs[0]
    r = p.add_run(f"[ SCREENSHOT: {screen_title} ]")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(WHITE)
    p2 = cell_title.add_paragraph()
    r2 = p2.add_run(f"URL: {url}")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(LIGHTGRAY)

    # Mockup area (large box - user can replace with real screenshot)
    set_cell_bg(cell_mock, LIGHT_BG)
    set_cell_border(cell_mock, LIGHTGRAY)
    cell_mock.height = Cm(8)
    # Instruction
    p = cell_mock.paragraphs[0]
    r = p.add_run(">> GANTI AREA INI DENGAN SCREENSHOT ASLI DARI nakesmart.com <<")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(MIDTRANS_BLUE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)

    # Mockup content (text-only representation)
    for line in mockup_lines:
        p = cell_mock.add_paragraph()
        line_type = line[0]
        text = line[1]
        if line_type == 'header':
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string(DEEP_FOREST)
            p.paragraph_format.space_after = Pt(2)
        elif line_type == 'card':
            r = p.add_run(f"  [{text}]")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string(CHARCOAL)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(1)
        elif line_type == 'button':
            r = p.add_run(f"  [ {text} ]")
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string(FOREST_GREEN)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
        elif line_type == 'price':
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor.from_string(FOREST_GREEN)
            p.paragraph_format.space_after = Pt(2)
        elif line_type == 'midtrans':
            r = p.add_run(f"[MIDTRANS BOX] {text}")
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor.from_string(MIDTRANS_BLUE)
            p.paragraph_format.space_after = Pt(2)
        elif line_type == 'success':
            r = p.add_run(f"[OK] PEMBAYARAN SUKSES - {text}")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string("15803D")
            p.paragraph_format.space_after = Pt(2)
        else:
            r = p.add_run(f"  {text}")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string(CHARCOAL)
            p.paragraph_format.space_after = Pt(1)

    # Caption below
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption.add_run(description)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.space_after = Pt(12)


def add_data_table(doc, headers, rows, col_widths_cm):
    """Add styled data table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell, FOREST_GREEN)
        set_cell_border(cell, FOREST_GREEN)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Body
    for ri, row in enumerate(rows):
        bg = LIGHT_BG if ri % 2 == 1 else WHITE
        for i, val in enumerate(row):
            cell = table.rows[ri + 1].cells[i]
            cell.width = Cm(col_widths_cm[i])
            set_cell_bg(cell, bg)
            set_cell_border(cell, LIGHTGRAY)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string(CHARCOAL)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # Spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_note_box(doc, text, color_bg="FEF3C7", color_border="F59E0B"):
    """Yellow note box."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, color_bg)
    set_cell_border(cell, color_border, size="12")
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(CHARCOAL)
    p.paragraph_format.space_after = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_docx():
    out_path = Path(__file__).parent / "Dokumen_Midtrans_Nakesmart.docx"
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ==================== COVER ====================
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NAKESMART")
    r.bold = True
    r.font.size = Pt(40)
    r.font.color.rgb = RGBColor.from_string(FOREST_GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("* * *")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(LIME)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DOKUMEN UNTUK MIDTRANS")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor.from_string(CHARCOAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("End-to-End Customer Journey + Payment Flow")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(DEEP_FOREST)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("+ Bukti Integrasi Zoom Partner")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(DEEP_FOREST)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DIAJUKAN UNTUK: PT MIDTRANS (Veritrans Indonesia)")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(FOREST_GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"EDISI {YEAR} - VERSION 1.0")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(FOREST_GREEN)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(COMPANY)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(CHARCOAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(WEBSITE)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(FOREST_GREEN)

    doc.add_page_break()

    # ==================== INSTRUKSI EDIT ====================
    add_heading(doc, "CARA EDIT DOKUMEN INI", level=1)
    add_para(doc,
        "Dokumen ini adalah versi DOCX editable dari Dokumen_Midtrans_Nakesmart.pdf. "
        "Setiap area screenshot ditandai dengan kotak biru bertuliskan "
        "\">> GANTI AREA INI DENGAN SCREENSHOT ASLI <<\".",
        size=10.5, align='justify')

    add_para(doc, "Langkah edit di Microsoft Word / Google Docs:", bold=True, size=10.5)
    add_bullet(doc, "Buka nakesmart.com sesuai URL pada setiap step")
    add_bullet(doc, "Screenshot tampilan layar (gunakan Snipping Tool / Windows + Shift + S / Mac Cmd+Shift+4)")
    add_bullet(doc, "Hapus kotak placeholder (klik kotak, tekan Delete)")
    add_bullet(doc, "Paste screenshot (Ctrl + V) di posisi yang sama")
    add_bullet(doc, "Resize gambar agar fit ke margin (drag pojok)")

    add_note_box(doc,
        "TIP: Untuk hasil terbaik, gunakan ukuran browser window 1280x720 atau 1366x768. "
        "Screenshot di mobile device juga diperbolehkan - tampilan responsive Nakesmart akan otomatis menyesuaikan.")

    doc.add_page_break()

    # ==================== SECTION 1: PROFIL ====================
    add_heading(doc, "1. PROFIL BISNIS NAKESMART", level=1)

    add_heading(doc, "1.1 Tentang Nakesmart", level=2)
    add_para(doc,
        "Nakesmart adalah Platform Ekosistem Kesehatan Digital Indonesia yang menghubungkan "
        "pasien dengan tenaga kesehatan (nakes), fasilitas kesehatan (faskes), dan event organizer "
        "kesehatan. Platform ini dioperasikan oleh PT GIANNA MEDICAL CENTER dengan domain "
        "resmi nakesmart.com dan beroperasi 24/7 untuk melayani masyarakat Indonesia.",
        align='justify')

    add_heading(doc, "1.2 Jenis Layanan yang Dijual", level=2)
    add_data_table(doc,
        headers=["Kategori Layanan", "Deskripsi", "Rentang Harga"],
        rows=[
            ["Telemedicine Reguler", "Konsultasi chat 30 menit dengan nakes terverifikasi", "Rp 45.000 - Rp 100.000"],
            ["Telemedicine Premium", "Konsultasi chat lengkap 45 menit (chat + voice + file)", "Rp 75.000 - Rp 200.000"],
            ["Telemedicine Exclusive", "Konsultasi video 15 menit via Zoom (partner Nakesmart)", "Rp 150.000 - Rp 500.000"],
            ["Booking Layanan Faskes", "Reservasi online layanan klinik/RS/lab", "Rp 50.000 - Rp 5.000.000"],
            ["Tiket Event Kesehatan", "Webinar, workshop, conference dari EO", "Rp 50.000 - Rp 5.000.000"],
            ["Top-up Wallet", "Pengisian saldo wallet untuk transaksi berikutnya", "Rp 10.000 - Rp 10.000.000"],
            ["Subscription Faskes", "Berlangganan SIMRS untuk fasilitas kesehatan", "Rp 499K - Rp 9.9jt/bulan"],
            ["Voucher & Paket", "Voucher diskon, bundle layanan, paket MCU", "Rp 50.000 - Rp 2.500.000"],
        ],
        col_widths_cm=[3.8, 8, 5],
    )

    add_heading(doc, "1.3 Customer Segments", level=2)
    add_bullet_rich(doc, [("Pasien (B2C)", True), (" - individu yang membutuhkan layanan kesehatan", False)])
    add_bullet_rich(doc, [("Nakes (B2B)", True), (" - tenaga kesehatan yang berpraktik di Nakesmart", False)])
    add_bullet_rich(doc, [("Faskes (B2B)", True), (" - klinik, rumah sakit, apotek, laboratorium", False)])
    add_bullet_rich(doc, [("EO (B2B)", True), (" - event organizer untuk event kesehatan", False)])

    doc.add_page_break()

    # ==================== SECTION 2: END-TO-END ====================
    add_heading(doc, "2. END-TO-END CUSTOMER JOURNEY", level=1)
    add_para(doc,
        "Berikut adalah alur lengkap perjalanan customer dari saat mengakses website hingga "
        "menerima layanan, dengan tampilan layar pada setiap tahap. Contoh kasus:",
        align='justify')
    add_para(doc, "Pasien membeli sesi Telemedicine Premium dengan dokter spesialis.",
             bold=True, color=FOREST_GREEN)

    # ============ STEPS DEFINITION ============
    steps = [
        {
            'num': 1,
            'title': 'CUSTOMER MEMBUKA WEBSITE NAKESMART',
            'desc': 'Customer membuka nakesmart.com di browser (desktop / mobile). '
                    'Halaman landing menampilkan layanan utama: Cari Faskes, Telemedicine, '
                    'Health Diary, AI Symptom Checker, dan Komunitas. '
                    'Customer login menggunakan email/Google OAuth.',
            'screen': 'Landing Page nakesmart.com',
            'url': 'https://nakesmart.com',
            'mockup': [
                ('header', 'Selamat Datang di NAKESMART'),
                ('text', 'Platform Kesehatan Digital Indonesia'),
                ('card', 'Telemedicine - Konsultasi online - Mulai Rp 45.000'),
                ('card', 'Cari Faskes - Klinik & RS terdekat - Booking online'),
                ('card', 'AI Symptom Checker - Cek gejala awal GRATIS'),
                ('button', 'MASUK / DAFTAR'),
            ],
            'caption': 'Tampilan halaman utama nakesmart.com (responsive untuk desktop & mobile)',
        },
        {
            'num': 2,
            'title': 'CUSTOMER MEMILIH LAYANAN TELEMEDICINE',
            'desc': 'Customer menelusuri katalog telemedicine. Dapat memfilter berdasarkan '
                    'spesialisasi (Dokter Umum, Anak, Penyakit Dalam, Kulit, dll), tier layanan '
                    '(Reguler / Premium / Exclusive), rating, dan harga.',
            'screen': 'Katalog Telemedicine',
            'url': 'nakesmart.com/telemedicine',
            'mockup': [
                ('header', 'Pilih Tier Konsultasi'),
                ('card', 'TIER REGULER (Chat) - Durasi 30 menit - Rp 45.000'),
                ('card', 'TIER PREMIUM (Chat+Voice) - Durasi 45 menit - Rp 75.000  [DIPILIH]'),
                ('card', 'TIER EXCLUSIVE (Video Zoom) - Durasi 15 menit - Rp 150.000'),
                ('text', 'dr. Aulia Rahmawati, Sp.PD'),
                ('text', 'Spesialis Penyakit Dalam - Rating 4.8'),
                ('button', 'PILIH NAKES INI'),
            ],
            'caption': 'Customer memilih Tier Premium - dr. Aulia Rahmawati, Sp.PD',
        },
        {
            'num': 3,
            'title': 'MEMILIH JADWAL & MASUK KE KERANJANG',
            'desc': 'Customer memilih slot waktu yang tersedia. Setelah memilih jadwal, '
                    'layanan otomatis masuk ke keranjang belanja. Customer dapat menambahkan '
                    'layanan lain atau langsung checkout.',
            'screen': 'Pilih Jadwal & Keranjang',
            'url': 'nakesmart.com/cart',
            'mockup': [
                ('header', 'Jadwal Konsultasi'),
                ('text', 'Tanggal: 12 Juni 2026'),
                ('text', 'Waktu  : 14:00 - 14:45 WIB'),
                ('header', 'Keranjang'),
                ('card', 'Telemedicine Premium - dr. Aulia Sp.PD - 45 menit - Rp 75.000'),
                ('text', 'Subtotal           : Rp 75.000'),
                ('text', 'Biaya Layanan      : Rp  2.500'),
                ('price', 'Total: Rp 77.500'),
                ('button', 'LANJUT KE CHECKOUT'),
            ],
            'caption': 'Sistem menghitung biaya layanan otomatis (PPN 0% untuk konsultasi kesehatan)',
        },
        {
            'num': 4,
            'title': 'CHECKOUT - PILIH METODE PEMBAYARAN MIDTRANS',
            'desc': 'Pada halaman checkout, customer memilih metode pembayaran. Nakesmart '
                    'mengintegrasikan Midtrans Snap sebagai payment gateway utama, yang '
                    'menyediakan multiple opsi pembayaran dalam satu tampilan '
                    '(Virtual Account, E-Wallet, QRIS, Kartu Kredit, dll).',
            'screen': 'Halaman Checkout',
            'url': 'nakesmart.com/checkout',
            'mockup': [
                ('header', 'Pilih Metode Pembayaran'),
                ('midtrans', 'Payment Gateway Resmi - 20+ metode pembayaran'),
                ('text', 'Saldo Wallet (Rp 25.000)'),
                ('text', 'Transfer Bank (Virtual Account)'),
                ('text', 'GoPay / OVO / DANA / ShopeePay'),
                ('text', 'QRIS'),
                ('text', 'Kartu Kredit / Debit'),
                ('price', 'Total: Rp 77.500'),
                ('button', 'BAYAR SEKARANG'),
            ],
            'caption': 'Customer memilih metode pembayaran lewat Midtrans Snap (popup)',
        },
        {
            'num': 5,
            'title': 'MIDTRANS SNAP - PROSES PEMBAYARAN',
            'desc': 'Setelah klik BAYAR SEKARANG, sistem Nakesmart memanggil API Midtrans Snap '
                    '/snap/v1/transactions untuk generate snap token. Popup Midtrans Snap '
                    'muncul dengan UI resmi Midtrans, menampilkan semua opsi pembayaran. '
                    'Customer memilih metode, lalu menyelesaikan pembayaran sesuai instruksi '
                    '(transfer ke VA, scan QRIS, dll).',
            'screen': 'Midtrans Snap Popup',
            'url': 'snap.midtrans.com/snap/v3/redirection/.../',
            'mockup': [
                ('midtrans', 'Order ID: NM-TELEMED-1734598123'),
                ('text', 'PT GIANNA MEDICAL CENTER'),
                ('price', 'Rp 77.500'),
                ('header', 'Pilih Pembayaran'),
                ('text', 'BCA Virtual Account'),
                ('text', 'BNI Virtual Account'),
                ('text', 'GoPay'),
                ('text', 'QRIS - All E-Wallet'),
                ('text', 'Kartu Kredit'),
                ('text', 'Powered by Midtrans (terdaftar di OJK)'),
            ],
            'caption': 'Popup Midtrans Snap - dikelola sepenuhnya oleh Midtrans (PCI-DSS compliant)',
        },
        {
            'num': 6,
            'title': 'CUSTOMER MENYELESAIKAN PEMBAYARAN',
            'desc': 'Customer menyelesaikan pembayaran via metode yang dipilih. '
                    'Contoh: GoPay - customer membuka aplikasi GoPay di HP, scan QR / konfirmasi '
                    'push notification, lalu memasukkan PIN. Pembayaran diproses oleh Midtrans '
                    '(real-time).',
            'screen': 'Konfirmasi GoPay (contoh)',
            'url': 'gopay.co.id/pay',
            'mockup': [
                ('header', 'Konfirmasi Pembayaran'),
                ('text', 'Untuk: PT Gianna Medical Center'),
                ('text', 'Via: Midtrans Payment Gateway'),
                ('price', 'Rp 77.500'),
                ('text', 'Saldo GoPay tersedia: Rp 350.000'),
                ('text', 'Masukkan PIN GoPay Anda:'),
                ('text', '* * * * * *'),
                ('button', 'KONFIRMASI BAYAR'),
            ],
            'caption': 'Alternatif: Transfer ke Virtual Account, scan QRIS, atau input kartu kredit',
        },
    ]

    for step in steps:
        add_step_header(doc, step['num'], step['title'])
        add_para(doc, step['desc'], align='justify')
        add_screenshot_placeholder(doc, step['screen'], step['url'], step['caption'], step['mockup'])

    # ============ STEP 7 - Webhook (table flow) ============
    add_step_header(doc, 7, "MIDTRANS MENGIRIM WEBHOOK NOTIFIKASI")
    add_para(doc,
        "Setelah pembayaran sukses, server Midtrans mengirim HTTP POST notification ke "
        "webhook endpoint Nakesmart di https://nakesmart.com/api/midtrans/notification. "
        "Server Nakesmart memverifikasi:",
        align='justify')
    add_bullet(doc, "Signature HMAC-SHA512 (signature_key) dari payload Midtrans untuk memastikan asli dari Midtrans")
    add_bullet(doc, "Gross_amount cocok dengan grand_total order di database")
    add_bullet(doc, "transaction_status = 'settlement' atau 'capture' dengan fraud_status = 'accept'")
    add_bullet(doc, "Idempotency check - menolak notifikasi duplikat untuk order yang sudah 'paid'")

    add_data_table(doc,
        headers=["CUSTOMER", "MIDTRANS", "NAKESMART BACKEND"],
        rows=[
            ["Konfirmasi PIN/OTP", "Process payment", "Wait for webhook"],
            ["",
             "POST webhook ke nakesmart.com/api/midtrans/notification",
             "Verify HMAC signature + Validate amount + Update status -> paid"],
            ["", "HTTP 200 OK", "Trigger fulfillment (book konsultasi)"],
        ],
        col_widths_cm=[5, 5.5, 6.5],
    )

    # ============ STEP 8 - Konfirmasi Sukses ============
    add_step_header(doc, 8, "KONFIRMASI PEMBAYARAN SUKSES KE CUSTOMER")
    add_para(doc,
        "Setelah backend Nakesmart memverifikasi pembayaran, customer otomatis di-redirect "
        "ke halaman sukses. Email konfirmasi dikirim ke alamat email customer. Konsultasi "
        "dengan dr. Aulia otomatis dijadwalkan untuk 12 Juni 2026 jam 14:00 WIB.",
        align='justify')
    add_screenshot_placeholder(doc, "Pembayaran Sukses",
        "nakesmart.com/payment/success",
        "Email + push notification dikirim, konsultasi siap dijadwalkan",
        [
            ('success', 'Order ID: NM-TELEMED-1734598123'),
            ('header', 'Detail Konsultasi'),
            ('text', 'Dokter : dr. Aulia Rahmawati, Sp.PD'),
            ('text', 'Tier   : Telemedicine Premium'),
            ('text', 'Jadwal : 12 Juni 2026, 14:00 WIB'),
            ('text', 'Durasi : 45 menit (Chat + Voice)'),
            ('price', 'Total Dibayar: Rp 77.500'),
            ('text', 'Metode: GoPay via Midtrans'),
            ('button', 'BUKA RUANG KONSULTASI'),
        ],
    )

    # ============ STEP 9 - Customer Menerima Layanan ============
    add_step_header(doc, 9, "CUSTOMER MENERIMA LAYANAN (FULFILLMENT)")
    add_para(doc,
        "Pada jadwal yang ditentukan (12 Juni 2026 14:00 WIB), customer membuka ruang "
        "konsultasi. Konsultasi berlangsung 45 menit dengan format chat + voice note. "
        "Setelah selesai, dokter dapat menerbitkan e-Resep digital yang otomatis tersimpan "
        "di akun customer.",
        align='justify')
    add_screenshot_placeholder(doc, "Ruang Konsultasi",
        "nakesmart.com/telemedicine/session/...",
        "Setelah sesi selesai, dana di-release ke wallet dokter (escrow protection)",
        [
            ('header', 'Konsultasi dengan dr. Aulia'),
            ('text', 'Sisa waktu: 38 menit | Premium - Chat+Voice'),
            ('card', 'dr. Aulia: Selamat sore Bapak/Ibu. Apa yang dapat saya bantu?'),
            ('card', 'Anda: Saya mengalami... (deskripsi gejala)'),
            ('card', 'dr. Aulia: Saya akan meresepkan... Silakan cek tab E-Resep'),
            ('text', 'Lihat E-Resep'),
            ('button', 'SELESAIKAN SESI'),
        ],
    )

    doc.add_page_break()

    # ==================== SECTION 3: TECHNICAL ====================
    add_heading(doc, "3. TECHNICAL FLOW - INTEGRASI MIDTRANS", level=1)
    add_para(doc,
        "Berikut adalah arsitektur teknis integrasi Midtrans pada platform Nakesmart, "
        "dari sisi backend dan frontend.",
        align='justify')

    add_heading(doc, "3.1 Arsitektur Tinggi", level=2)
    add_data_table(doc,
        headers=["Layer", "Komponen", "Fungsi"],
        rows=[
            ["Frontend", "Web App (React) + Mobile App (Expo)",
             "Trigger checkout, embed Snap.js, redirect ke status page"],
            ["Backend API", "FastAPI di nakesmart.com/api",
             "Generate snap token, handle webhook, update DB"],
            ["Database", "Supabase (PostgreSQL)",
             "Store orders, payments, escrow status"],
            ["Payment Gateway", "Midtrans Snap (snap.midtrans.com)",
             "Process payment dari customer, kirim webhook ke backend"],
            ["Notification", "Midtrans webhook -> Backend",
             "POST notification ke /api/midtrans/notification"],
        ],
        col_widths_cm=[3.5, 5.5, 8],
    )

    add_heading(doc, "3.2 Endpoint Webhook Midtrans di Nakesmart", level=2)
    add_bullet_rich(doc, [("URL Notifikasi (Notification URL): ", True),
                          ("https://nakesmart.com/api/midtrans/notification", False)])
    add_bullet_rich(doc, [("Finish Redirect URL: ", True),
                          ("https://nakesmart.com/payment/success", False)])
    add_bullet_rich(doc, [("Unfinish URL: ", True),
                          ("https://nakesmart.com/payment/pending", False)])
    add_bullet_rich(doc, [("Error URL: ", True),
                          ("https://nakesmart.com/payment/failed", False)])

    add_heading(doc, "3.3 Keamanan & Verifikasi", level=2)
    add_bullet_rich(doc, [("HMAC-SHA512 Signature", True),
                          (" - Setiap webhook diverifikasi signature_key sesuai dokumentasi Midtrans", False)])
    add_bullet_rich(doc, [("Amount Validation", True),
                          (" - gross_amount divalidasi sama dengan grand_total order", False)])
    add_bullet_rich(doc, [("Idempotency", True),
                          (" - Order yang sudah berstatus 'paid' tidak diproses ulang", False)])
    add_bullet_rich(doc, [("HTTPS Only", True),
                          (" - Semua komunikasi menggunakan TLS 1.2+", False)])
    add_bullet_rich(doc, [("Anti-Replay", True),
                          (" - Order ID unik, tidak bisa digunakan ulang", False)])
    add_bullet_rich(doc, [("PCI-DSS Compliance", True),
                          (" - Tidak ada data kartu kredit yang disimpan di server Nakesmart (semua handled oleh Midtrans)", False)])

    add_heading(doc, "3.4 Refund Flow", level=2)
    add_para(doc,
        "Apabila customer cancel layanan atau dispute disetujui, Nakesmart memproses refund "
        "melalui salah satu jalur berikut: (1) Refund ke wallet internal customer untuk "
        "digunakan transaksi berikutnya, atau (2) Refund langsung ke rekening sumber pembayaran "
        "via Midtrans Refund API.",
        align='justify')

    doc.add_page_break()

    # ==================== SECTION 4: ZOOM PARTNER ====================
    add_heading(doc, "4. BUKTI INTEGRASI ZOOM PARTNER", level=1)
    add_para(doc,
        "Nakesmart adalah partner resmi Zoom Video Communications, Inc. melalui Zoom Marketplace "
        "App. Integrasi ini digunakan untuk layanan Telemedicine Tier Exclusive yang membutuhkan "
        "video konsultasi HD antara pasien dan nakes.",
        align='justify')

    add_heading(doc, "4.1 Detail Integrasi", level=2)
    add_data_table(doc,
        headers=["Item", "Detail"],
        rows=[
            ["Partner Type", "Zoom Marketplace - Server-to-Server OAuth App"],
            ["App Name", "Nakesmart Telemedicine Video"],
            ["Account ID", "[ZOOM_ACCOUNT_ID] (dikonfigurasi via env variable)"],
            ["Client ID", "[ZOOM_CLIENT_ID] (dikonfigurasi via env variable)"],
            ["Scopes", "meeting:write:admin, meeting:read:admin, recording:read:admin, user:read:admin"],
            ["Webhook URL", "https://nakesmart.com/api/zoom/webhook"],
            ["Use Case", "Auto-create video meeting untuk Telemedicine Tier Exclusive"],
            ["Pricing Tier", "Zoom Business / Pro (multi-host)"],
            ["Capacity", "500 meetings concurrent (upgradable)"],
            ["Recording", "Cloud recording enabled (with patient consent)"],
            ["Code Reference", "backend/services/zoom_api.py, backend/services/zoom_webinar_service.py"],
        ],
        col_widths_cm=[4, 13],
    )

    add_heading(doc, "4.2 Tampilan Video Meeting (Telemedicine Exclusive)", level=2)
    add_screenshot_placeholder(doc, "Video Konsultasi Zoom",
        "nakesmart.com/telemedicine/zoom/...",
        "Tampilan video meeting di Telemedicine Exclusive (powered by Zoom)",
        [
            ('header', 'Sesi Video Konsultasi'),
            ('text', 'Powered by Zoom Video SDK'),
            ('card', '[VIDEO FEED] - dr. Aulia Rahmawati, Sp.PD - HD Video aktif'),
            ('card', '[VIDEO FEED ANDA] - Patient camera aktif - Mic: ON'),
            ('text', 'Mute Mic'),
            ('text', 'Share Screen'),
            ('button', 'AKHIRI SESI'),
        ],
    )

    add_heading(doc, "4.3 Verifikasi Partnership Zoom", level=2)
    add_para(doc, "Untuk verifikasi keaslian integrasi, Midtrans dapat memeriksa hal berikut:")

    add_bullet_rich(doc, [("Zoom Marketplace Listing", True),
                          (" - https://marketplace.zoom.us (Search: 'Nakesmart Telemedicine')", False)])
    add_bullet_rich(doc, [("Account Dashboard", True),
                          (" - Login ke marketplace.zoom.us dengan akun terdaftar Nakesmart", False)])
    add_bullet_rich(doc, [("App Type", True),
                          (" - Server-to-Server OAuth (untuk auto-create meetings tanpa user login)", False)])
    add_bullet_rich(doc, [("Verifikasi via API", True),
                          (" - Endpoint Zoom https://api.zoom.us/v2/users/me dapat dipanggil dengan credentials Nakesmart", False)])
    add_bullet_rich(doc, [("Recording Storage", True),
                          (" - Recording tersimpan di Zoom Cloud (sesuai retention policy)", False)])
    add_bullet_rich(doc, [("HIPAA-Style Compliance", True),
                          (" - Setting privacy untuk konsultasi kesehatan (waiting room, password meeting, encrypted recording)", False)])

    add_note_box(doc,
        "CATATAN UNTUK MIDTRANS: Kredensial Zoom (Account ID, Client ID, Client Secret) tidak "
        "dilampirkan di dokumen ini demi keamanan. Apabila diperlukan, Nakesmart dapat memberikan "
        "screenshot dashboard Zoom Marketplace yang menampilkan App yang telah terinstal, melalui "
        "channel komunikasi yang aman (email terenkripsi atau in-person).")

    doc.add_page_break()

    # ==================== SECTION 5: COMPLIANCE ====================
    add_heading(doc, "5. KEPATUHAN & LEGAL", level=1)

    add_heading(doc, "5.1 Pengembalian Dana (Refund Policy)", level=2)
    add_para(doc,
        "Nakesmart memiliki kebijakan pengembalian dana yang transparan dan terdokumentasi di "
        "nakesmart.com/refund-policy. Skenario refund yang umum:",
        align='justify')
    add_bullet_rich(doc, [("Nakes tidak hadir", True),
                          (" - Refund 100% otomatis ke wallet (sistem auto-detect setelah 30 menit)", False)])
    add_bullet_rich(doc, [("Customer cancel H-7", True), (" - Refund 100%", False)])
    add_bullet_rich(doc, [("Customer cancel H-3 sampai H-1", True), (" - Refund 50%", False)])
    add_bullet_rich(doc, [("Customer cancel < 3 jam", True), (" - Tidak ada refund", False)])
    add_bullet_rich(doc, [("Dispute disetujui admin", True), (" - Refund 100% via Midtrans Refund API", False)])
    add_bullet_rich(doc, [("Customer Service ticket", True), (" - Customer dapat mengajukan komplain via tab Bantuan", False)])

    add_heading(doc, "5.2 Penyelesaian Sengketa (Dispute Resolution)", level=2)
    add_para(doc,
        "Customer yang ingin mengajukan komplain dapat menghubungi: "
        "(1) Customer Service via in-app chat (24/7), "
        "(2) Email: support@nakesmart.com (response < 24 jam), "
        "(3) WhatsApp business: +62 851-xxxx-xxxx. "
        "Admin Nakesmart memiliki kewenangan untuk approve refund maksimal Rp 5.000.000 per ticket. "
        "Apabila tidak ada penyelesaian, customer dapat mengeskalasi ke YLKI atau OJK.",
        align='justify')

    add_heading(doc, "5.3 Data Privacy (UU PDP)", level=2)
    add_para(doc,
        "Nakesmart patuh pada UU No. 27 Tahun 2022 tentang Pelindungan Data Pribadi (UU PDP). "
        "Data customer (termasuk data kesehatan yang sensitif) disimpan di server Indonesia "
        "(Supabase ap-southeast region), dienkripsi at-rest dan in-transit. Detail privacy "
        "policy: nakesmart.com/privacy-policy.",
        align='justify')

    add_heading(doc, "5.4 Kategori Bisnis (MCC)", level=2)
    add_para(doc,
        "Merchant Category Code: 8099 (Health Services - Other), 8062 (Hospitals), "
        "5912 (Drug Stores & Pharmacies). Nakesmart bertransaksi sebagai marketplace healthcare "
        "yang menghubungkan pasien, nakes, faskes, dan event organizer.",
        align='justify')

    doc.add_page_break()

    # ==================== SECTION 6: RINGKASAN & KONTAK ====================
    add_heading(doc, "6. RINGKASAN & KONTAK", level=1)

    add_heading(doc, "6.1 Ringkasan Customer Journey", level=2)
    add_data_table(doc,
        headers=["Step", "Aksi", "Sistem", "Waktu"],
        rows=[
            ["1", "Buka nakesmart.com", "Web Nakesmart", "< 5 detik"],
            ["2", "Pilih layanan (telemedicine)", "Web Nakesmart", "1-3 menit"],
            ["3", "Pilih jadwal & add to cart", "Web Nakesmart", "< 1 menit"],
            ["4", "Checkout & pilih metode bayar", "Web + Midtrans", "< 1 menit"],
            ["5", "Buka Midtrans Snap popup", "Midtrans", "< 5 detik"],
            ["6", "Customer bayar via GoPay/VA/QRIS", "Midtrans + Bank/E-Wallet", "1-10 menit"],
            ["7", "Webhook Midtrans -> Backend Nakesmart", "Server-to-Server", "< 5 detik"],
            ["8", "Konfirmasi sukses ke customer", "Email + Push + Web", "Instant"],
            ["9", "Konsultasi pada jadwal yang dipesan", "Web/Mobile + Zoom (Exclusive)", "15-45 menit"],
        ],
        col_widths_cm=[1.2, 6.3, 5.5, 4],
    )

    add_heading(doc, "6.2 Kontak PIC Nakesmart untuk Midtrans", level=2)
    add_bullet_rich(doc, [("PT GIANNA MEDICAL CENTER", True)])
    add_bullet_rich(doc, [("Domain Resmi: ", True), ("nakesmart.com", False)])
    add_bullet_rich(doc, [("Email Sales/Billing: ", True), ("finance@nakesmart.com", False)])
    add_bullet_rich(doc, [("Email Teknis/Integrasi: ", True), ("tech@nakesmart.com", False)])
    add_bullet_rich(doc, [("Customer Service: ", True), ("support@nakesmart.com", False)])
    add_bullet_rich(doc, [("WhatsApp Business: ", True), ("+62 851-xxxx-xxxx", False)])
    add_bullet_rich(doc, [("Alamat Kantor: ", True), ("Jakarta, Indonesia", False)])

    doc.add_paragraph()

    # Closing note box
    closing_table = doc.add_table(rows=1, cols=1)
    cell = closing_table.rows[0].cells[0]
    set_cell_bg(cell, LIGHT_BG)
    set_cell_border(cell, FOREST_GREEN, size="16")
    p = cell.paragraphs[0]
    r = p.add_run(
        "Demikian dokumen end-to-end customer journey dan bukti integrasi Zoom partner "
        "Nakesmart kami sampaikan untuk keperluan onboarding Midtrans Production."
    )
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(FOREST_GREEN)
    p.paragraph_format.space_after = Pt(8)

    p2 = cell.add_paragraph()
    r2 = p2.add_run(
        "Apabila diperlukan dokumen pendukung tambahan (NIB, NPWP, akta pendirian, "
        "screenshot dashboard Zoom Marketplace, atau testing transaksi di sandbox), "
        "silakan menghubungi PIC kami via email tech@nakesmart.com."
    )
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(CHARCOAL)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(f"Jakarta, {YEAR}")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(CHARCOAL)

    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(COMPANY)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(FOREST_GREEN)

    p = doc.add_paragraph()
    r = p.add_run(WEBSITE)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(CHARCOAL)

    # Footer
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f"© {YEAR} {COMPANY}  -  {WEBSITE}  -  Dokumen Midtrans")
    fr.font.size = Pt(8.5)
    fr.font.color.rgb = RGBColor.from_string(CHARCOAL)

    doc.save(str(out_path))
    print(f"[OK] Dokumen Midtrans DOCX dibuat: {out_path}")
    print(f"     Ukuran: {out_path.stat().st_size // 1024} KB")


if __name__ == '__main__':
    build_docx()
